import os
import json
from flask import Blueprint, render_template, request, send_file, jsonify
from docx import Document
import pdfkit

bp_arbolp = Blueprint('arbolp', __name__, template_folder='templates')

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
RUTA_ARBOL = os.path.join(DATA_DIR, "arbol_problemas.json")
os.makedirs(DATA_DIR, exist_ok=True)
WKHTMLTOPDF_PATH = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"  # Ajusta si tu ruta es diferente
PDFKIT_CONFIG = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH)


# ------------------ Rutas ------------------
@bp_arbolp.route('/arbolp')
def ruta_arbolp():
    municipio = request.args.get('municipio', '')
    nodos = {}
    if os.path.exists(RUTA_ARBOL):
        with open(RUTA_ARBOL, 'r', encoding='utf-8') as f:
            data = json.load(f)
            nodos = data.get(municipio, {})
    return render_template('arbolp.html', municipio=municipio, nodos=nodos)

@bp_arbolp.route('/guardar_arbol', methods=['POST'])
def guardar_arbol():
    datos = request.json
    municipio = datos.get('municipio')
    if not municipio:
        return jsonify({"status": "error", "msg": "Municipio no definido"}), 400

    if os.path.exists(RUTA_ARBOL):
        with open(RUTA_ARBOL, 'r', encoding='utf-8') as f:
            data = json.load(f)
    else:
        data = {}

    data[municipio] = {
        "titulo": datos.get("titulo", "Problema central"),
        "efectos": datos.get("efectos", []),
        "causas": datos.get("causas", [])
    }

    with open(RUTA_ARBOL, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4, ensure_ascii=False)

    return jsonify({"status": "ok"})

# ------------------ Exportar PDF ------------------
@bp_arbolp.route('/exportar_pdf', methods=['POST'])
def exportar_pdf():
    municipio = request.form.get("municipio")
    if not os.path.exists(RUTA_ARBOL):
        with open(RUTA_ARBOL, 'w', encoding='utf-8') as f:
            json.dump({}, f)

    with open(RUTA_ARBOL, 'r', encoding='utf-8') as f:
        data = json.load(f)

    arbol = data.get(municipio, {"titulo": "Problema central", "efectos": [], "causas": []})

    # Generar HTML para el PDF (igual que tu app)
    html = f"""
    <html>
    <head>
      <style>
        body {{ font-family: Arial, sans-serif; text-align: center; }}
        .nivel {{ margin: 20px 0; }}
        .efectos {{ color: green; }}
        .problema {{ font-weight: bold; font-size: 1.2em; }}
        .causas {{ color: blue; }}
      </style>
    </head>
    <body>
      <h1>🌳 Árbol de Problemas - {municipio}</h1>
      <div class="nivel efectos">
        <h2>Efectos</h2>
        {"".join(f"<p>- {e['titulo']}</p>" for e in arbol.get('efectos', []))}
      </div>
      <div class="nivel problema">
        <h2>Problema central</h2>
        <p>{arbol.get('titulo')}</p>
      </div>
      <div class="nivel causas">
        <h2>Causas</h2>
        {"".join(f"<p>- {c['titulo']}</p>" for c in arbol.get('causas', []))}
      </div>
    </body>
    </html>
    """

    pdf_path = os.path.join(DATA_DIR, f"arbol_problemas_{municipio}.pdf")
    pdfkit.from_string(html, pdf_path, configuration=PDFKIT_CONFIG)
    return send_file(pdf_path, as_attachment=True)

# ------------------ Exportar Word ------------------
@bp_arbolp.route('/exportar_word', methods=['POST'])
def exportar_word():
    municipio = request.form.get("municipio")
    if not os.path.exists(RUTA_ARBOL):
        with open(RUTA_ARBOL, 'w', encoding='utf-8') as f:
            json.dump({}, f)

    with open(RUTA_ARBOL, 'r', encoding='utf-8') as f:
        data = json.load(f)

    arbol = data.get(municipio, {"titulo": "Problema central", "efectos": [], "causas": []})

    docx_path = os.path.join(DATA_DIR, f"arbol_problemas_{municipio}.docx")
    doc = Document()
    doc.add_heading("🌳 Árbol de Problemas", 0)
    doc.add_heading(f"Problema central: {arbol['titulo']}", level=2)

    if arbol.get("efectos"):
        doc.add_heading("Efectos:", level=3)
        for e in arbol["efectos"]:
            doc.add_paragraph(f"- {e['titulo']}")
    if arbol.get("causas"):
        doc.add_heading("Causas:", level=3)
        for c in arbol["causas"]:
            doc.add_paragraph(f"- {c['titulo']}")

    doc.save(docx_path)
    return send_file(docx_path, as_attachment=True)
